#!/usr/bin/env python3
"""Validate DOCX structure, grounding, accessibility, privacy, and visual QA."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from collections import Counter
from contextlib import ExitStack
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

SKILL_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SKILL_DIR))

from lib.package_safety import admitted_docx  # noqa: E402
from lib.worddoc import (  # noqa: E402
    audit_request_grounding,
    inspect_docx,
    iter_all_paragraphs,
    load_json,
    paragraph_heading_level,
    validate_json,
)
from lib.path_safety import atomic_write_text  # noqa: E402


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--request", type=Path)
    parser.add_argument("--visual-qa", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--overwrite-report", action="store_true")
    parser.add_argument("--final", action="store_true")
    parser.add_argument("--privacy", action="store_true")
    parser.add_argument(
        "--allow-privacy-findings",
        action="store_true",
        help="Allow custom properties or possible PII during final privacy validation.",
    )
    parser.add_argument("--strict-warnings", action="store_true")
    parser.add_argument("--allow-pending", action="store_true")
    parser.add_argument(
        "--include-content",
        action="store_true",
        help="Include document text and detailed paths in protected diagnostics.",
    )
    parser.add_argument(
        "--trust-render-isolation-claim",
        action="store_true",
        help=(
            "Acknowledge that network isolation is producer-asserted metadata, "
            "not an independently authenticated attestation."
        ),
    )
    return parser.parse_args()


def issue(code: str, message: str, **details: Any) -> dict[str, Any]:
    value: dict[str, Any] = {"code": code, "message": message}
    value.update(details)
    return value


def normalize_heading(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip()).casefold()


def _resolve_from(base: Path, value: str) -> Path:
    result = Path(value)
    return result.resolve() if result.is_absolute() else (base / result).resolve()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def validate_visual_qa(
    path: Path,
    document_identity: Path,
    document_content: Path,
    *,
    include_content: bool,
    trust_isolation_claim: bool,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    data = load_json(path)
    validate_json(data, SKILL_DIR / "schemas" / "visual-qa.schema.json")
    errors: list[dict[str, Any]] = []
    document_digest = sha256_file(document_content)
    qa_document = _resolve_from(path.parent, data["document"])
    if qa_document != document_identity.resolve():
        errors.append(
            issue(
                "visual-document-mismatch",
                "Visual QA refers to a different document",
                **(
                    {
                        "expected": str(document_identity.resolve()),
                        "actual": str(qa_document),
                    }
                    if include_content
                    else {}
                ),
            )
        )

    render_report_path = _resolve_from(path.parent, data["render_report"])
    if not render_report_path.is_file():
        errors.append(
            issue(
                "missing-render-report",
                "Visual QA render report does not exist",
                **(
                    {"path": str(render_report_path)}
                    if include_content
                    else {}
                ),
            )
        )
        render_report = {}
    else:
        render_report = load_json(render_report_path)
        rendered_document_value = render_report.get("document")
        if rendered_document_value is not None and _resolve_from(
            render_report_path.parent, str(rendered_document_value)
        ) != document_identity.resolve():
            errors.append(
                issue(
                    "render-document-mismatch",
                    "Render report refers to a different document",
                    **(
                        {
                            "expected": str(document_identity.resolve()),
                            "actual": str(rendered_document_value),
                        }
                        if include_content
                        else {}
                    ),
                )
            )
        if render_report.get("pages") != data["rendered_pages"]:
            errors.append(
                issue(
                    "render-page-count-mismatch",
                    "Visual QA page count differs from render report",
                    qa_pages=data["rendered_pages"],
                    report_pages=render_report.get("pages"),
                )
            )
        if render_report.get("document_sha256") != document_digest:
            errors.append(
                issue(
                    "render-document-digest-mismatch",
                    "Render report is not bound to the current document content",
                )
            )
        if render_report.get("network_isolation") is not True:
            errors.append(
                issue(
                    "render-network-isolation-missing",
                    "Final visual QA requires a network-isolated render",
                )
            )
        elif not trust_isolation_claim:
            errors.append(
                issue(
                    "render-isolation-claim-untrusted",
                    "Producer-asserted isolation requires explicit trust acknowledgement",
                )
            )
        page_images = render_report.get("page_images")
        if not isinstance(page_images, list) or len(page_images) != data["rendered_pages"]:
            errors.append(
                issue(
                    "rendered-page-images-incomplete",
                    "Render report must list exactly one image for every rendered page",
                )
            )
            page_images = []
        elif len(
            {
                _resolve_from(render_report_path.parent, value)
                for value in page_images
                if isinstance(value, str)
            }
        ) != len(page_images):
            errors.append(
                issue(
                    "duplicate-rendered-page-images",
                    "Each rendered page must have its own image",
                )
            )
        missing_images = [
            value
            for value in page_images
            if not isinstance(value, str)
            or not _resolve_from(render_report_path.parent, value).is_file()
        ]
        if missing_images:
            errors.append(
                issue(
                    "missing-rendered-page-images",
                    "Rendered page images listed in the report are missing",
                    count=len(missing_images),
                )
            )
        image_paths = [
            _resolve_from(render_report_path.parent, value)
            for value in page_images
            if isinstance(value, str)
            and _resolve_from(render_report_path.parent, value).is_file()
        ]
        image_digests = render_report.get("page_image_sha256")
        expected_image_names = {image.name for image in image_paths}
        if (
            not isinstance(image_digests, dict)
            or set(image_digests) != expected_image_names
            or len(expected_image_names) != len(image_paths)
        ):
            errors.append(
                issue(
                    "rendered-page-digests-incomplete",
                    "Render report must bind every page image to a SHA-256 digest",
                )
            )
        else:
            tampered_images = sum(
                image_digests.get(image.name) != sha256_file(image)
                for image in image_paths
            )
            if tampered_images:
                errors.append(
                    issue(
                        "rendered-page-digest-mismatch",
                        "Rendered page content differs from the render report",
                        count=tampered_images,
                    )
                )

        pdf_value = render_report.get("pdf")
        pdf_path = (
            _resolve_from(render_report_path.parent, pdf_value)
            if isinstance(pdf_value, str)
            else None
        )
        if pdf_path is None or not pdf_path.is_file():
            errors.append(
                issue(
                    "missing-rendered-pdf",
                    "Render report must reference the rendered PDF",
                )
            )
        elif render_report.get("pdf_sha256") != sha256_file(pdf_path):
            errors.append(
                issue(
                    "rendered-pdf-digest-mismatch",
                    "Rendered PDF content differs from the render report",
                )
            )

    expected = set(range(1, data["rendered_pages"] + 1))
    inspected = set(data["inspected_pages"])
    if inspected != expected:
        errors.append(
            issue(
                "visual-pages-incomplete",
                "Every rendered page must be inspected",
                missing=sorted(expected - inspected),
                unexpected=sorted(inspected - expected),
            )
        )
    blocking = [
        finding
        for finding in data["findings"]
        if finding["severity"] == "blocking" and not finding["resolved"]
    ]
    if blocking:
        errors.append(
            issue(
                "visual-blocking-findings",
                "Blocking visual findings remain unresolved",
                count=len(blocking),
            )
        )
    if data["status"] != "passed":
        errors.append(
            issue("visual-qa-not-passed", "Visual QA status must be passed")
        )
    return data, errors


def main() -> int:
    args = parse_args()
    admissions = ExitStack()
    errors: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    visual_qa = None
    unverified_claims: list[str] = []
    try:
        source_path = args.document.absolute()
        if not source_path.is_file():
            raise FileNotFoundError(f"Document not found: {source_path}")
        source = admissions.enter_context(admitted_docx(source_path))
        inspection = inspect_docx(source)
        document = Document(str(source))

        titles = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style and paragraph.style.name == "Title" and paragraph.text.strip()
        ]
        if not titles:
            errors.append(issue("missing-title", "No non-empty Title paragraph found"))

        previous_level = 0
        for heading in inspection["headings"]:
            level = heading["level"]
            if previous_level and level > previous_level + 1:
                errors.append(
                    issue(
                        "heading-level-skip",
                        f"Heading level jumps from {previous_level} to {level}",
                        **(
                            {"heading": heading["text"]}
                            if args.include_content
                            else {}
                        ),
                    )
                )
            previous_level = level

        paragraph_texts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if len(paragraph.text.strip()) >= 40
        ]
        duplicates = [text for text, count in Counter(paragraph_texts).items() if count > 1]
        if duplicates:
            warnings.append(
                issue(
                    "duplicate-paragraphs",
                    "Long duplicate paragraphs detected",
                    count=len(duplicates),
                    **(
                        {"samples": duplicates[:5]}
                        if args.include_content
                        else {}
                    ),
                )
            )

        empty_spacers = 0
        direct_runs = 0
        total_runs = 0
        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                has_semantic_content = bool(
                    paragraph._p.xpath(".//w:drawing | .//w:br | .//w:fldChar")
                )
                if not has_semantic_content:
                    empty_spacers += 1
        for paragraph in iter_all_paragraphs(document):
            for run in paragraph.runs:
                total_runs += 1
                properties = run._r.find(qn("w:rPr"))
                if properties is not None and len(properties):
                    direct_runs += 1
        if empty_spacers:
            warnings.append(
                issue(
                    "empty-spacing-paragraphs",
                    "Empty paragraphs may be used for visual spacing",
                    count=empty_spacers,
                )
            )
        direct_ratio = round(direct_runs / total_runs, 4) if total_runs else 0
        if direct_ratio > 0.25:
            warnings.append(
                issue(
                    "excess-direct-formatting",
                    "More than 25% of runs contain direct formatting",
                    ratio=direct_ratio,
                )
            )

        if inspection["tables"] != inspection["tables_with_repeating_header"]:
            errors.append(
                issue(
                    "table-header-not-repeated",
                    "Every table must mark its first row as a repeating header",
                    tables=inspection["tables"],
                    compliant=inspection["tables_with_repeating_header"],
                )
            )
        if inspection["inline_figures"] != inspection["figures_with_alt_text"]:
            errors.append(
                issue(
                    "missing-figure-alt-text",
                    "Every inline figure must have alternative text",
                    figures=inspection["inline_figures"],
                    compliant=inspection["figures_with_alt_text"],
                )
            )

        placeholders = inspection["placeholders"]
        accidental = {
            key: value
            for key, value in placeholders.items()
            if key != "pending-confirmation"
        }
        if accidental:
            errors.append(
                issue(
                    "accidental-placeholders",
                    "Unresolved TODO, XXX, or template variables remain",
                    placeholders=accidental,
                )
            )
        if placeholders.get("pending-confirmation") and not args.allow_pending:
            errors.append(
                issue(
                    "pending-confirmation",
                    "Pending confirmations remain in the document",
                    count=placeholders["pending-confirmation"],
                )
            )

        if args.request:
            request = load_json(args.request.resolve())
            validate_json(
                request, SKILL_DIR / "schemas" / "document-request.schema.json"
            )
            required = {
                normalize_heading(value)
                for value in request.get("required_sections", [])
            }
            actual = {
                normalize_heading(item["text"]) for item in inspection["headings"]
            }
            missing = sorted(required - actual)
            if missing:
                errors.append(
                    issue(
                        "missing-required-sections",
                        "Required sections are absent",
                        count=len(missing),
                        **(
                            {"sections": missing}
                            if args.include_content
                            else {}
                        ),
                    )
                )
            grounding = audit_request_grounding(request)
            if grounding["source_reference_issues"]:
                errors.append(
                    issue(
                        "source-reference-issues",
                        "Source-backed blocks have missing or unknown source IDs",
                        count=len(grounding["source_reference_issues"]),
                        **(
                            {
                                "details": grounding[
                                    "source_reference_issues"
                                ]
                            }
                            if args.include_content
                            else {}
                        ),
                    )
                )
            if grounding["unclassified_blocks"]:
                warnings.append(
                    issue(
                        "unclassified-content",
                        "Content blocks lack provenance classification",
                        count=len(grounding["unclassified_blocks"]),
                        **(
                            {"details": grounding["unclassified_blocks"]}
                            if args.include_content
                            else {}
                        ),
                    )
                )
            unverified_claims = grounding["unverified_claims"]
            if unverified_claims and not args.allow_pending:
                errors.append(
                    issue(
                        "unverified-claims",
                        "Unverified claims remain",
                        count=len(unverified_claims),
                        **(
                            {"claims": unverified_claims}
                            if args.include_content
                            else {}
                        ),
                    )
                )

        if args.final:
            package = inspection["package"]
            if package["comments"]:
                errors.append(
                    issue("comments-remain", "Comments remain in the final document")
                )
            if package["insertions"] or package["deletions"]:
                errors.append(
                    issue(
                        "tracked-changes-remain",
                        "Tracked insertions or deletions remain in the final document",
                        insertions=package["insertions"],
                        deletions=package["deletions"],
                    )
                )

        if args.privacy:
            properties = inspection["core_properties"]
            personal = {
                key: properties[key]
                for key in (
                    "author",
                    "comments",
                    "created",
                    "keywords",
                    "last_modified_by",
                    "modified",
                )
                if properties.get(key)
            }
            if personal:
                errors.append(
                    issue(
                        "personal-metadata",
                        "Personal author metadata remains",
                        properties=sorted(personal),
                    )
                )
            if inspection["package"]["custom_properties"]:
                privacy_issue = issue(
                    "custom-properties",
                    "Custom document properties remain",
                )
                if args.final and not args.allow_privacy_findings:
                    errors.append(privacy_issue)
                else:
                    warnings.append(privacy_issue)
            if inspection["package"]["extended_personal_metadata"]:
                errors.append(
                    issue(
                        "extended-personal-metadata",
                        "Extended document properties contain private metadata",
                        fields=inspection["package"][
                            "extended_personal_metadata"
                        ],
                    )
                )
            if inspection["package"]["comment_personal_metadata"]:
                errors.append(
                    issue(
                        "comment-personal-metadata",
                        "Comment metadata contains personal fields",
                        fields=inspection["package"][
                            "comment_personal_metadata"
                        ],
                    )
                )
            zip_metadata = inspection["package"]["zip_metadata"]
            if any(zip_metadata.values()):
                errors.append(
                    issue(
                        "zip-metadata",
                        "ZIP container metadata has not been normalized",
                        fields=sorted(
                            key for key, value in zip_metadata.items() if value
                        ),
                    )
                )
            if inspection["possible_pii"]:
                privacy_issue = issue(
                    "possible-pii",
                    "Possible personal data appears in document content",
                    counts=inspection["possible_pii"],
                )
                if args.final and not args.allow_privacy_findings:
                    errors.append(privacy_issue)
                else:
                    warnings.append(privacy_issue)

        if args.visual_qa:
            visual_qa, visual_errors = validate_visual_qa(
                args.visual_qa.resolve(),
                source_path,
                source,
                include_content=args.include_content,
                trust_isolation_claim=args.trust_render_isolation_claim,
            )
            errors.extend(visual_errors)
        elif args.final:
            errors.append(
                issue(
                    "missing-visual-qa",
                    "Final validation requires a visual QA report",
                )
            )

        valid = not errors and not (args.strict_warnings and warnings)
        report = {
            "document": str(source_path),
            "valid": valid,
            "errors": errors,
            "warnings": warnings,
            "metrics": {
                "paragraphs": inspection["paragraphs"],
                "headings": len(inspection["headings"]),
                "tables": inspection["tables"],
                "figures": inspection["inline_figures"],
                "sections": len(inspection["sections"]),
                "direct_formatting_ratio": direct_ratio,
                "empty_spacing_paragraphs": empty_spacers,
            },
            "visual_qa": (
                visual_qa
                if args.include_content or visual_qa is None
                else {
                    "status": visual_qa.get("status"),
                    "rendered_pages": visual_qa.get("rendered_pages"),
                    "inspected_page_count": len(
                        visual_qa.get("inspected_pages", [])
                    ),
                    "finding_count": len(visual_qa.get("findings", [])),
                }
            ),
            "unverified_claims": (
                unverified_claims if args.include_content else []
            ),
            "unverified_claim_count": len(unverified_claims),
        }
        validate_json(
            report, SKILL_DIR / "schemas" / "validation-report.schema.json"
        )
        if args.report:
            atomic_write_text(
                args.report,
                json.dumps(report, indent=2, ensure_ascii=False) + "\n",
                overwrite=args.overwrite_report,
                label="validation report",
            )
        print(json.dumps(report, indent=2, ensure_ascii=False))
        return 0 if valid else 1
    except Exception as exc:
        print(
            json.dumps(
                {
                    "document": str(args.document),
                    "valid": False,
                    "errors": [
                        issue(
                            "validation-exception",
                            (
                                str(exc)
                                if args.include_content
                                else "Validation failed; rerun with "
                                "--include-content in a protected workspace"
                            ),
                            exception_type=type(exc).__name__,
                        )
                    ],
                    "warnings": warnings,
                    "metrics": {},
                    "visual_qa": visual_qa,
                    "unverified_claims": unverified_claims,
                },
                ensure_ascii=False,
            )
        )
        return 1
    finally:
        admissions.close()


if __name__ == "__main__":
    raise SystemExit(main())
