#!/usr/bin/env python3
"""End-to-end fixture tests for krt-word-illuminator scripts."""

from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from hashlib import sha256
from pathlib import Path
from unittest.mock import patch
from xml.etree import ElementTree as ET

from docx import Document

ROOT = Path(__file__).resolve().parent
SKILL_DIR = ROOT.parent
sys.path.insert(0, str(SKILL_DIR))

from lib import path_safety  # noqa: E402


def run_script(
    name: str, *args: str, env: dict[str, str] | None = None
) -> tuple[int, dict]:
    completed = subprocess.run(
        [sys.executable, str(ROOT / name), *args],
        check=False,
        capture_output=True,
        text=True,
        env=env,
    )
    try:
        payload = json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise AssertionError(
            f"{name} emitted invalid JSON\nstdout={completed.stdout}\nstderr={completed.stderr}"
        ) from exc
    return completed.returncode, payload


def rewrite_docx(path: Path, replacements: dict[str, bytes]) -> None:
    temporary = path.with_suffix(".rewrite.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as target:
        existing = set(source.namelist())
        for info in source.infolist():
            target.writestr(info, replacements.get(info.filename, source.read(info)))
        for name, value in replacements.items():
            if name not in existing:
                target.writestr(name, value)
    temporary.replace(path)


class WordIlluminatorTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(
            prefix="krt-word-illuminator-test-"
        )
        self.work = Path(self.temporary.name)
        self.request = self.work / "request.json"
        self.output = self.work / "report.docx"
        self.request.write_text(
            json.dumps(
                {
                    "document_type": "technical_report",
                    "objective": "Document the verified system architecture",
                    "title": "Architecture Report",
                    "subtitle": "Verified baseline",
                    "audience": "Engineering and management",
                    "language": "en",
                    "output": "report.docx",
                    "required_sections": ["Executive summary", "Architecture"],
                    "sources": [{"id": "requirements", "path": "requirements.md"}],
                    "document": {
                        "toc": True,
                        "header": "Architecture Report",
                        "footer": "Internal",
                        "page_numbers": True
                    },
                    "sections": [
                        {
                            "heading": "Executive summary",
                            "level": 1,
                            "blocks": [
                                {
                                    "type": "paragraph",
                                    "text": "The system uses a documented service boundary.",
                                    "provenance": "source",
                                    "source_ids": ["requirements"]
                                }
                            ]
                        },
                        {
                            "heading": "Architecture",
                            "level": 1,
                            "blocks": [
                                {
                                    "type": "table",
                                    "headers": ["Component", "Responsibility"],
                                    "rows": [["API", "Request handling"]],
                                    "caption": "Table 1. Components",
                                    "provenance": "source",
                                    "source_ids": ["requirements"]
                                },
                                {
                                    "type": "note",
                                    "text": "Implementation detail supplied by the user.",
                                    "provenance": "user"
                                }
                            ]
                        }
                    ],
                    "unverified_claims": []
                }
            ),
            encoding="utf-8",
        )

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def create_document(self) -> dict:
        code, result = run_script("create_docx.py", str(self.request))
        self.assertEqual(code, 0, result)
        self.assertTrue(self.output.is_file())
        return result

    def visual_qa(self, document: Path | None = None) -> Path:
        document = document or self.output
        prefix = "" if document == self.output else f"{document.stem}-"
        page_image = self.work / f"{prefix}page-1.png"
        page_image.write_bytes(b"fixture")
        rendered_pdf = self.work / f"{prefix}report.pdf"
        rendered_pdf.write_bytes(b"%PDF-1.4 fixture")
        render_report = self.work / f"{prefix}render-report.json"
        render_report.write_text(
            json.dumps(
                {
                    "document": str(document),
                    "document_sha256": sha256(document.read_bytes()).hexdigest(),
                    "network_isolation": True,
                    "pdf": str(rendered_pdf),
                    "pdf_sha256": sha256(rendered_pdf.read_bytes()).hexdigest(),
                    "pages": 1,
                    "page_images": [str(page_image)],
                    "page_image_sha256": {
                        page_image.name: sha256(page_image.read_bytes()).hexdigest()
                    },
                }
            ),
            encoding="utf-8",
        )
        path = self.work / f"{prefix}visual-qa.json"
        path.write_text(
            json.dumps(
                {
                    "document": str(document),
                    "render_report": str(render_report),
                    "rendered_pages": 1,
                    "inspected_pages": [1],
                    "status": "passed",
                    "findings": []
                }
            ),
            encoding="utf-8",
        )
        return path

    def test_create_inspect_and_validate(self) -> None:
        result = self.create_document()
        self.assertEqual(result["tables"], 1)
        self.assertEqual(result["visual_qa"], "pending_render_and_manual_inspection")

        code, inspection = run_script(
            "inspect_docx.py", str(self.output), "--include-content"
        )
        self.assertEqual(code, 0, inspection)
        self.assertEqual(inspection["tables"], 1)
        self.assertEqual(inspection["tables_with_repeating_header"], 1)
        self.assertEqual(
            [item["text"] for item in inspection["headings"]],
            ["Contents", "Executive summary", "Architecture"],
        )

        code, validation = run_script(
            "validate_docx.py",
            str(self.output),
            "--request",
            str(self.request),
        )
        self.assertEqual(code, 0, validation)
        self.assertTrue(validation["valid"])

        code, final_validation = run_script(
            "validate_docx.py",
            str(self.output),
            "--request",
            str(self.request),
            "--visual-qa",
            str(self.visual_qa()),
            "--final",
            "--trust-render-isolation-claim",
        )
        self.assertEqual(code, 0, final_validation)

    def test_edit_and_compare(self) -> None:
        self.create_document()
        patch = self.work / "patch.json"
        edited = self.work / "edited.docx"
        patch.write_text(
            json.dumps(
                {
                    "operations": [
                        {
                            "op": "replace_paragraph",
                            "old": "Implementation detail supplied by the user.",
                            "new": "Updated implementation detail."
                        },
                        {
                            "op": "insert_after_heading",
                            "heading": "Architecture",
                            "blocks": [
                                {
                                    "type": "paragraph",
                                    "text": "Inserted scope statement.",
                                    "provenance": "user"
                                }
                            ]
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        code, edit = run_script(
            "edit_docx.py",
            str(self.output),
            str(patch),
            "--output",
            str(edited),
        )
        self.assertEqual(code, 0, edit)
        code, comparison = run_script(
            "compare_docx.py",
            str(self.output),
            str(edited),
            "--include-content",
        )
        self.assertEqual(code, 0, comparison)
        self.assertTrue(comparison["changed"])
        self.assertTrue(comparison["text_diff"])

    def test_default_inspection_redacts_content_and_metadata_values(self) -> None:
        self.create_document()
        document = Document(str(self.output))
        document.core_properties.author = "Sentinel Private Author"
        document.paragraphs[-1].text = "Sentinel Private Content"
        document.save(self.output)

        code, inspection = run_script("inspect_docx.py", str(self.output))

        self.assertEqual(code, 0, inspection)
        serialized = json.dumps(inspection)
        self.assertNotIn("Sentinel Private Author", serialized)
        self.assertNotIn("Sentinel Private Content", serialized)
        self.assertIn(
            "author", inspection["core_properties"]["populated_fields"]
        )

    def test_edit_preserves_single_run_formatting(self) -> None:
        source = self.work / "formatted.docx"
        document = Document()
        run = document.add_paragraph().add_run("Original")
        run.bold = True
        document.save(source)
        patch = self.work / "formatted-patch.json"
        patch.write_text(
            json.dumps(
                {
                    "operations": [
                        {
                            "op": "replace_paragraph",
                            "old": "Original",
                            "new": "Updated",
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        output = self.work / "formatted-edited.docx"
        code, result = run_script(
            "edit_docx.py", str(source), str(patch), "--output", str(output)
        )
        self.assertEqual(code, 0, result)
        paragraph = Document(str(output)).paragraphs[0]
        self.assertEqual(paragraph.text, "Updated")
        self.assertTrue(paragraph.runs[0].bold)

    def test_edit_rejects_complex_paragraph_replacement(self) -> None:
        source = self.work / "complex.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Bold ").bold = True
        paragraph.add_run("text")
        document.save(source)
        patch = self.work / "complex-patch.json"
        patch.write_text(
            json.dumps(
                {
                    "operations": [
                        {
                            "op": "replace_paragraph",
                            "old": "Bold text",
                            "new": "Updated",
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        output = self.work / "complex-edited.docx"
        code, result = run_script(
            "edit_docx.py", str(source), str(patch), "--output", str(output)
        )
        self.assertNotEqual(code, 0)
        self.assertIn("simple single-run text paragraphs", result["error"])
        self.assertFalse(output.exists())

    def test_privacy_scrub(self) -> None:
        self.create_document()
        document = Document(str(self.output))
        document.core_properties.author = "Sensitive Author"
        document.core_properties.last_modified_by = "Sensitive Editor"
        commented = document.paragraphs[-1]
        document.add_comment(
            commented.runs,
            text="Working comment",
            author="Sensitive Reviewer",
            initials="SR",
        )
        document.save(self.output)
        clean = self.work / "clean.docx"
        code, result = run_script(
            "privacy_scrub.py",
            str(self.output),
            "--output",
            str(clean),
            "--remove-comments",
        )
        self.assertEqual(code, 0, result)
        self.assertEqual(result["before"]["comments"], 1)
        self.assertEqual(result["after"]["comments"], 0)
        self.assertIn(
            "author", result["before"]["comment_personal_metadata"]
        )
        self.assertEqual(result["after"]["comment_personal_metadata"], [])
        self.assertIn(
            "author", result["before"]["populated_personal_metadata_fields"]
        )
        self.assertIn(
            "last_modified_by",
            result["before"]["populated_personal_metadata_fields"],
        )
        self.assertEqual(result["after"]["populated_personal_metadata_fields"], [])
        self.assertEqual(result["after"]["extended_personal_metadata"], [])
        self.assertFalse(any(result["after"]["zip_metadata"].values()))
        self.assertNotIn("Sensitive Author", json.dumps(result))
        self.assertNotIn("Sensitive Editor", json.dumps(result))
        self.assertTrue(clean.is_file())

    def test_privacy_scan_handles_alternate_comment_prefix(self) -> None:
        self.create_document()
        document = Document(str(self.output))
        paragraph = document.paragraphs[-1]
        document.add_comment(
            paragraph.runs,
            text="Contact prefix-sentinel@example.invalid",
            author="Prefix Sentinel",
            initials="PS",
        )
        document.save(self.output)
        with zipfile.ZipFile(self.output) as archive:
            comments = archive.read("word/comments.xml")
        alternate_prefix = (
            comments.replace(b"xmlns:w=", b"xmlns:x=")
            .replace(b"<w:", b"<x:")
            .replace(b"</w:", b"</x:")
            .replace(b" w:", b" x:")
        )
        rewrite_docx(
            self.output,
            {"word/comments.xml": alternate_prefix},
        )

        code, result = run_script(
            "inspect_docx.py",
            str(self.output),
            "--include-content",
        )

        self.assertEqual(code, 0, result)
        self.assertEqual(result["package"]["comments"], 1)
        self.assertIn("author", result["package"]["comment_personal_metadata"])
        self.assertEqual(result["possible_pii"]["email"], 1)

    def test_unknown_source_id_fails_creation(self) -> None:
        request = json.loads(self.request.read_text(encoding="utf-8"))
        request["sections"][0]["blocks"][0]["source_ids"] = ["missing"]
        self.request.write_text(json.dumps(request), encoding="utf-8")
        code, result = run_script("create_docx.py", str(self.request))
        self.assertNotEqual(code, 0)
        self.assertIn("Source grounding failed", result["error"])
        self.assertFalse(self.output.exists())

    def test_request_output_cannot_escape_request_directory(self) -> None:
        request = json.loads(self.request.read_text(encoding="utf-8"))
        escaped = self.work.parent / f"{self.work.name}-escaped.docx"
        request["output"] = f"../{escaped.name}"
        self.request.write_text(json.dumps(request), encoding="utf-8")

        code, result = run_script("create_docx.py", str(self.request))

        self.assertNotEqual(code, 0)
        self.assertIn("escapes the approved input roots", result["error"])
        self.assertFalse(escaped.exists())

    def test_output_symlink_is_rejected_even_with_overwrite(self) -> None:
        target = self.work / "target.docx"
        target.write_bytes(b"must remain unchanged")
        link = self.work / "linked.docx"
        link.symlink_to(target)

        code, result = run_script(
            "create_docx.py",
            str(self.request),
            "--output",
            str(link),
            "--overwrite",
        )

        self.assertNotEqual(code, 0)
        self.assertIn("symbolic-link component", result["error"])
        self.assertEqual(target.read_bytes(), b"must remain unchanged")

    def test_public_entry_points_reject_source_symlinks(self) -> None:
        self.create_document()
        link = self.work / "source-link.docx"
        link.symlink_to(self.output)
        patch_file = self.work / "symlink-patch.json"
        patch_file.write_text(
            json.dumps(
                {
                    "operations": [
                        {
                            "op": "replace_paragraph",
                            "old": "Implementation detail supplied by the user.",
                            "new": "Updated",
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        binary_dir = self.work / "symlink-bin"
        binary_dir.mkdir()
        for name in ("libreoffice", "pdftoppm"):
            executable = binary_dir / name
            executable.write_text("#!/bin/sh\nexit 1\n", encoding="utf-8")
            executable.chmod(0o755)
        environment = dict(os.environ)
        environment["PATH"] = (
            f"{binary_dir}{os.pathsep}{environment['PATH']}"
        )
        commands = (
            ("inspect_docx.py", (str(link),), None),
            (
                "edit_docx.py",
                (
                    str(link),
                    str(patch_file),
                    "--output",
                    str(self.work / "symlink-edited.docx"),
                ),
                None,
            ),
            ("compare_docx.py", (str(link), str(self.output)), None),
            (
                "privacy_scrub.py",
                (
                    str(link),
                    "--output",
                    str(self.work / "symlink-clean.docx"),
                ),
                None,
            ),
            (
                "render_docx.py",
                (
                    str(link),
                    "--output-dir",
                    str(self.work / "symlink-render"),
                    "--allow-networked-render",
                ),
                environment,
            ),
            ("validate_docx.py", (str(link), "--include-content"), None),
        )

        for script, arguments, script_environment in commands:
            with self.subTest(script=script):
                code, result = run_script(
                    script,
                    *arguments,
                    env=script_environment,
                )
                self.assertNotEqual(code, 0, result)
                self.assertIn("source-open-failed", json.dumps(result))

    def test_create_rejects_symlinked_template(self) -> None:
        self.create_document()
        template_link = self.work / "template-link.docx"
        template_link.symlink_to(self.output)
        request = json.loads(self.request.read_text(encoding="utf-8"))
        request["template"] = template_link.name
        request["output"] = "from-linked-template.docx"
        linked_request = self.work / "linked-template-request.json"
        linked_request.write_text(json.dumps(request), encoding="utf-8")

        code, result = run_script("create_docx.py", str(linked_request))

        self.assertNotEqual(code, 0, result)
        self.assertIn("symbolic-link component", result["error"])
        self.assertFalse((self.work / "from-linked-template.docx").exists())

    def test_creation_report_is_no_clobber_by_default(self) -> None:
        report = self.output.with_suffix(".creation-report.json")
        report.write_text("preserve", encoding="utf-8")

        code, result = run_script("create_docx.py", str(self.request))

        self.assertNotEqual(code, 0)
        self.assertIn("Refusing to overwrite existing report", result["error"])
        self.assertEqual(report.read_text(encoding="utf-8"), "preserve")
        self.assertFalse(self.output.exists())

    def test_multi_artifact_publish_restores_prior_files_on_failure(self) -> None:
        first_output = self.work / "first.txt"
        second_output = self.work / "second.txt"
        first_output.write_text("old-first", encoding="utf-8")
        second_output.write_text("old-second", encoding="utf-8")
        first_prepared = self.work / ".first.prepared"
        second_prepared = self.work / ".second.prepared"
        first_prepared.write_text("new-first", encoding="utf-8")
        second_prepared.write_text("new-second", encoding="utf-8")
        real_publish = path_safety.atomic_publish_file

        def fail_second(prepared: Path, output: Path, **kwargs: object) -> Path:
            if output == second_output:
                raise OSError("injected second-artifact failure")
            return real_publish(prepared, output, **kwargs)

        with patch(
            "lib.path_safety.atomic_publish_file",
            side_effect=fail_second,
        ):
            with self.assertRaisesRegex(OSError, "second-artifact"):
                path_safety.atomic_publish_files(
                    (
                        (first_prepared, first_output, "first artifact"),
                        (second_prepared, second_output, "second artifact"),
                    ),
                    overwrite=True,
                )

        self.assertEqual(first_output.read_text(encoding="utf-8"), "old-first")
        self.assertEqual(second_output.read_text(encoding="utf-8"), "old-second")

    def test_final_validation_rejects_incomplete_render_evidence(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        report_path = self.work / "render-report.json"
        report = json.loads(report_path.read_text(encoding="utf-8"))
        report["page_images"] = []
        report_path.write_text(json.dumps(report), encoding="utf-8")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "rendered-page-images-incomplete",
            [item["code"] for item in result["errors"]],
        )

    def test_validation_redacts_document_content_by_default(self) -> None:
        self.create_document()
        document = Document(str(self.output))
        sentinel = "Sentinel confidential paragraph repeated for diagnostics"
        document.add_paragraph(sentinel)
        document.add_paragraph(sentinel)
        document.save(self.output)

        code, result = run_script("validate_docx.py", str(self.output))

        self.assertEqual(code, 0, result)
        self.assertNotIn(sentinel, json.dumps(result))
        self.assertIn(
            "duplicate-paragraphs",
            [item["code"] for item in result["warnings"]],
        )

    def test_inspection_report_is_no_clobber_and_rejects_symlink(self) -> None:
        self.create_document()
        report = self.work / "inspection.json"
        report.write_text("preserve", encoding="utf-8")

        code, result = run_script(
            "inspect_docx.py",
            str(self.output),
            "--report",
            str(report),
        )
        self.assertNotEqual(code, 0)
        self.assertIn("Refusing to overwrite", result["error"])
        self.assertEqual(report.read_text(encoding="utf-8"), "preserve")

        target = self.work / "inspection-target.json"
        target.write_text("target", encoding="utf-8")
        link = self.work / "inspection-link.json"
        link.symlink_to(target)
        code, result = run_script(
            "inspect_docx.py",
            str(self.output),
            "--report",
            str(link),
            "--overwrite",
        )
        self.assertNotEqual(code, 0)
        self.assertIn("symbolic-link component", result["error"])
        self.assertEqual(target.read_text(encoding="utf-8"), "target")

    def test_final_validation_rejects_stale_render_document_digest(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        document = Document(str(self.output))
        document.add_paragraph("Post-render change")
        document.save(self.output)

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "render-document-digest-mismatch",
            [item["code"] for item in result["errors"]],
        )

    def test_final_validation_rejects_networked_render(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        render_report = self.work / "render-report.json"
        report = json.loads(render_report.read_text(encoding="utf-8"))
        report["network_isolation"] = False
        render_report.write_text(json.dumps(report), encoding="utf-8")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "render-network-isolation-missing",
            [item["code"] for item in result["errors"]],
        )

    def test_final_validation_requires_explicit_trust_for_isolation_claim(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        render_report = self.work / "render-report.json"
        report = json.loads(render_report.read_text(encoding="utf-8"))
        report["network_isolation_evidence"] = "independently-authenticated"
        render_report.write_text(json.dumps(report), encoding="utf-8")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )

        self.assertNotEqual(code, 0)
        self.assertIn(
            "render-isolation-claim-untrusted",
            [item["code"] for item in result["errors"]],
        )

    def test_final_validation_rejects_tampered_page_image(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        (self.work / "page-1.png").write_bytes(b"changed after render")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "rendered-page-digest-mismatch",
            [item["code"] for item in result["errors"]],
        )

    def test_final_validation_rejects_incomplete_page_digests(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        render_report = self.work / "render-report.json"
        report = json.loads(render_report.read_text(encoding="utf-8"))
        report["page_image_sha256"] = {}
        render_report.write_text(json.dumps(report), encoding="utf-8")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "rendered-page-digests-incomplete",
            [item["code"] for item in result["errors"]],
        )

    def test_final_validation_rejects_tampered_pdf(self) -> None:
        self.create_document()
        visual_qa = self.visual_qa()
        (self.work / "report.pdf").write_bytes(b"%PDF-1.4 changed")

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "rendered-pdf-digest-mismatch",
            [item["code"] for item in result["errors"]],
        )

    def test_final_privacy_blocks_possible_pii_without_explicit_exception(self) -> None:
        self.create_document()
        document = Document(str(self.output))
        document.add_paragraph("Contact: person@example.com")
        document.save(self.output)
        visual_qa = self.visual_qa()

        code, result = run_script(
            "validate_docx.py",
            str(self.output),
            "--visual-qa",
            str(visual_qa),
            "--final",
            "--privacy",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "possible-pii",
            [item["code"] for item in result["errors"]],
        )

    def test_privacy_scrub_cleans_metadata_and_detects_pii_in_footnotes(self) -> None:
        self.create_document()
        with zipfile.ZipFile(self.output) as archive:
            app_root = ET.fromstring(archive.read("docProps/app.xml"))
        company = next(
            node
            for node in app_root.iter()
            if node.tag.rsplit("}", 1)[-1] == "Company"
        )
        company.text = "Sentinel Private Company"
        footnotes = (
            b"<w:footnotes xmlns:w='http://schemas.openxmlformats.org/"
            b"wordprocessingml/2006/main'><w:footnote w:id='1'><w:p><w:r>"
            b"<w:t>private-footnote@example.com</w:t></w:r></w:p>"
            b"</w:footnote></w:footnotes>"
        )
        rewrite_docx(
            self.output,
            {
                "docProps/app.xml": ET.tostring(
                    app_root, encoding="utf-8", xml_declaration=True
                ),
                "word/footnotes.xml": footnotes,
            },
        )
        clean = self.work / "metadata-clean.docx"
        code, scrub = run_script(
            "privacy_scrub.py",
            str(self.output),
            "--output",
            str(clean),
        )
        self.assertEqual(code, 0, scrub)
        self.assertIn(
            "Company", scrub["before"]["extended_personal_metadata"]
        )
        self.assertEqual(scrub["after"]["extended_personal_metadata"], [])
        self.assertGreater(scrub["after"]["possible_pii"].get("email", 0), 0)
        self.assertFalse(any(scrub["after"]["zip_metadata"].values()))

        code, validation = run_script(
            "validate_docx.py",
            str(clean),
            "--visual-qa",
            str(self.visual_qa(clean)),
            "--final",
            "--privacy",
        )
        self.assertNotEqual(code, 0)
        self.assertIn(
            "possible-pii",
            [item["code"] for item in validation["errors"]],
        )

    def test_render_pipeline_records_pending_visual_qa(self) -> None:
        self.create_document()
        binary_dir = self.work / "bin"
        binary_dir.mkdir()
        libreoffice = binary_dir / "libreoffice"
        libreoffice.write_text(
            """#!/usr/bin/env python3
import sys
from pathlib import Path
output = Path(sys.argv[sys.argv.index("--outdir") + 1])
source = Path(sys.argv[-1])
(output / f"{source.stem}.pdf").write_bytes(b"%PDF-1.4 fixture")
""",
            encoding="utf-8",
        )
        pdftoppm = binary_dir / "pdftoppm"
        pdftoppm.write_text(
            """#!/usr/bin/env python3
import sys
from pathlib import Path
from PIL import Image
prefix = Path(sys.argv[-1])
Image.new("RGB", (100, 140), "white").save(f"{prefix}-1.png")
""",
            encoding="utf-8",
        )
        libreoffice.chmod(0o755)
        pdftoppm.chmod(0o755)
        environment = dict(os.environ)
        environment["PATH"] = f"{binary_dir}{os.pathsep}{environment['PATH']}"
        render_dir = self.work / "render"
        code, result = run_script(
            "render_docx.py",
            str(self.output),
            "--output-dir",
            str(render_dir),
            "--allow-networked-render",
            env=environment,
        )
        self.assertEqual(code, 0, result)
        self.assertEqual(result["pages"], 1)
        self.assertEqual(result["visual_qa"], "pending_manual_inspection")
        self.assertIsInstance(result["network_isolation"], bool)
        self.assertIn(
            result["network_isolation_method"],
            {"unshare-network-namespace", "explicit-networked-preview"},
        )
        self.assertTrue((render_dir / "page-1.png").is_file())
        self.assertTrue((render_dir / "render-report.json").is_file())
        sentinel = render_dir / "unmanaged-user-artifact.txt"
        sentinel.write_text("preserve", encoding="utf-8")

        code, rejected = run_script(
            "render_docx.py",
            str(self.output),
            "--output-dir",
            str(render_dir),
            "--overwrite",
            "--allow-networked-render",
            env=environment,
        )

        self.assertNotEqual(code, 0, rejected)
        self.assertIn("unknown or unsafe entries", rejected["error"])
        self.assertEqual(sentinel.read_text(encoding="utf-8"), "preserve")
        self.assertTrue((render_dir / "render-report.json").is_file())
        sentinel.unlink()

        code, replaced = run_script(
            "render_docx.py",
            str(self.output),
            "--output-dir",
            str(render_dir),
            "--overwrite",
            "--allow-networked-render",
            env=environment,
        )
        self.assertEqual(code, 0, replaced)
        self.assertTrue((render_dir / "render-report.json").is_file())

    def test_failed_render_overwrite_preserves_previous_evidence(self) -> None:
        self.create_document()
        binary_dir = self.work / "failing-bin"
        binary_dir.mkdir()
        libreoffice = binary_dir / "libreoffice"
        libreoffice.write_text("#!/bin/sh\nexit 9\n", encoding="utf-8")
        libreoffice.chmod(0o755)
        environment = dict(os.environ)
        environment["PATH"] = (
            f"{binary_dir}{os.pathsep}{environment['PATH']}"
        )
        render_dir = self.work / "render"
        render_dir.mkdir()
        sentinel = render_dir / "prior-evidence.txt"
        sentinel.write_text("preserve", encoding="utf-8")

        code, result = run_script(
            "render_docx.py",
            str(self.output),
            "--output-dir",
            str(render_dir),
            "--overwrite",
            "--allow-networked-render",
            env=environment,
        )

        self.assertNotEqual(code, 0, result)
        self.assertEqual(sentinel.read_text(encoding="utf-8"), "preserve")
        self.assertEqual(
            sorted(path.name for path in render_dir.iterdir()),
            ["prior-evidence.txt"],
        )

    def test_bundled_template_preserves_single_page_field(self) -> None:
        template = (
            ROOT.parent / "assets" / "templates" / "professional-report.docx"
        ).resolve()
        self.assertTrue(template.is_file())
        request = json.loads(self.request.read_text(encoding="utf-8"))
        request["template"] = str(template)
        request["output"] = "templated.docx"
        request["document"] = {"toc": False}
        self.request.write_text(json.dumps(request), encoding="utf-8")
        output = self.work / "templated.docx"
        code, result = run_script("create_docx.py", str(self.request))
        self.assertEqual(code, 0, result)
        code, inspection = run_script("inspect_docx.py", str(output))
        self.assertEqual(code, 0, inspection)
        self.assertEqual(inspection["package"]["fields"], 1)

    def test_build_template_is_no_clobber_by_default(self) -> None:
        output = self.work / "template.docx"
        output.write_bytes(b"preserve")

        code, result = run_script(
            "build_template.py",
            "--output",
            str(output),
        )

        self.assertNotEqual(code, 0, result)
        self.assertEqual(output.read_bytes(), b"preserve")

        code, result = run_script(
            "build_template.py",
            "--output",
            str(output),
            "--overwrite",
        )
        self.assertEqual(code, 0, result)
        self.assertTrue(zipfile.is_zipfile(output))

    def test_ambiguous_edit_fails(self) -> None:
        document = Document()
        document.add_paragraph("Repeated")
        document.add_paragraph("Repeated")
        source = self.work / "ambiguous.docx"
        document.save(source)
        patch = self.work / "ambiguous.json"
        patch.write_text(
            json.dumps(
                {
                    "operations": [
                        {
                            "op": "replace_paragraph",
                            "old": "Repeated",
                            "new": "Changed"
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        code, result = run_script(
            "edit_docx.py",
            str(source),
            str(patch),
            "--output",
            str(self.work / "ambiguous-edited.docx"),
        )
        self.assertNotEqual(code, 0)
        self.assertIn("Ambiguous paragraph replacement", result["error"])


if __name__ == "__main__":
    unittest.main()
