"""Shared DOCX construction and inspection helpers."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

from lib.package_safety import admitted_docx

SKILL_DIR = Path(__file__).resolve().parents[1]
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

PLACEHOLDER_PATTERNS = {
    "TODO": re.compile(r"\bTODO\b", re.IGNORECASE),
    "XXX": re.compile(r"\bXXX\b", re.IGNORECASE),
    "template-variable": re.compile(r"\{\{[^{}]+\}\}"),
    "pending-confirmation": re.compile(
        r"\[(?:PENDING CONFIRMATION|PENDIENTE DE CONFIRMACI[ÓO]N)\]",
        re.IGNORECASE,
    ),
}

PII_PATTERNS = {
    "email": re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE),
    "phone": re.compile(r"(?<!\w)(?:\+?\d[\d .()/-]{7,}\d)(?!\w)"),
}


def load_json(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError(f"Expected a JSON object in {path}")
    return data


def validate_json(data: dict[str, Any], schema_path: Path) -> None:
    try:
        import jsonschema
    except ImportError as exc:  # pragma: no cover - environment failure
        raise RuntimeError("jsonschema is required to validate request files") from exc
    schema = load_json(schema_path)
    try:
        jsonschema.validate(data, schema)
    except jsonschema.ValidationError as error:
        location = "/".join(str(item) for item in error.absolute_path)
        suffix = f" at {location}" if location else ""
        raise ValueError(f"JSON does not conform to schema{suffix}") from error


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def clear_document_body(document: DocumentType) -> None:
    """Remove template body content while preserving final section properties."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        properties.append(marker)
    marker.set(qn("w:val"), "true")


def _set_row_cant_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        properties.append(marker)


def _set_keep_next(paragraph: Any, enabled: bool = True) -> None:
    properties = paragraph._p.get_or_add_pPr()
    marker = properties.find(qn("w:keepNext"))
    if enabled and marker is None:
        properties.append(OxmlElement("w:keepNext"))
    elif not enabled and marker is not None:
        properties.remove(marker)


def _set_alt_text(inline_shape: Any, alt_text: str) -> None:
    properties = inline_shape._inline.xpath(".//wp:docPr")
    if not properties:
        properties = inline_shape._inline.xpath(".//pic:cNvPr")
    for node in properties:
        node.set("descr", alt_text)
        if not node.get("title"):
            node.set("title", alt_text[:80])


def _add_field(paragraph: Any, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_node, separate])
    if placeholder:
        run._r.append(OxmlElement("w:t"))
        run._r[-1].text = placeholder
    run._r.append(end)


def _has_field(container: Any, instruction: str) -> bool:
    wanted = instruction.strip().casefold()
    for node in container._element.xpath(".//w:instrText"):
        if (node.text or "").strip().casefold().startswith(wanted):
            return True
    return False


def ensure_styles(document: DocumentType, normalize_builtins: bool = False) -> None:
    styles = document.styles
    if normalize_builtins:
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.08

        styles["Title"].font.name = "Aptos Display"
        styles["Title"].font.size = Pt(30)
        for level, size in ((1, 18), (2, 14), (3, 12)):
            heading = styles[f"Heading {level}"]
            heading.font.name = "Aptos Display"
            heading.font.size = Pt(size)
            heading.font.color.rgb = RGBColor(31, 78, 121)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(5)
            heading.paragraph_format.keep_with_next = True

    custom = {
        "Code": ("Consolas", 9, "F3F4F6"),
        "Note": ("Aptos", 10, "EAF2F8"),
        "Warning": ("Aptos", 10, "FDEDEC"),
    }
    for name, (font_name, size, _) in custom.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            style.font.name = font_name
            style.font.size = Pt(size)
            style.paragraph_format.left_indent = Cm(0.5)
            style.paragraph_format.right_indent = Cm(0.5)
            style.paragraph_format.space_before = Pt(4)
            style.paragraph_format.space_after = Pt(6)


def configure_document(
    document: DocumentType,
    options: dict[str, Any],
    *,
    preserve_template_layout: bool,
) -> None:
    page_size = options.get("page_size")
    margins = options.get("margins_cm", {})
    for section in document.sections:
        if page_size == "A4" or (not page_size and not preserve_template_layout):
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        elif page_size == "LETTER":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        if not preserve_template_layout or margins:
            section.top_margin = Cm(float(margins.get("top", 2.2)))
            section.right_margin = Cm(float(margins.get("right", 2.0)))
            section.bottom_margin = Cm(float(margins.get("bottom", 2.2)))
            section.left_margin = Cm(float(margins.get("left", 2.0)))

        if "header" in options:
            header = section.header.paragraphs[0]
            header.text = options["header"]
            header.style = document.styles["Normal"]
        if "footer" in options:
            footer = section.footer.paragraphs[0]
            footer.text = options["footer"]
            footer.style = document.styles["Normal"]
        if options.get("page_numbers", True) and not _has_field(section.footer, "PAGE"):
            footer = section.footer.paragraphs[0]
            if footer.text:
                footer.add_run("  ·  ")
            _add_field(footer, "PAGE", "1")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_core_properties(document: DocumentType, values: dict[str, Any]) -> None:
    properties = document.core_properties
    mapping = {
        "title": "title",
        "subject": "subject",
        "author": "author",
        "keywords": "keywords",
        "comments": "comments",
    }
    for key, attribute in mapping.items():
        if key in values:
            setattr(properties, attribute, str(values[key]))


def _source_ids(request: dict[str, Any]) -> set[str]:
    ids: set[str] = set()
    for source in request.get("sources", []):
        if isinstance(source, str):
            ids.add(Path(source).stem)
        else:
            ids.add(str(source["id"]))
    return ids


def resolve_request_path(
    value: str,
    base_dir: Path,
    *,
    label: str,
    trusted_roots: Iterable[Path] = (),
) -> Path:
    """Resolve a data-supplied path without allowing absolute or parent escapes."""
    supplied = Path(value)
    root = base_dir.resolve()
    lexical = (supplied if supplied.is_absolute() else root / supplied).absolute()
    resolved = lexical.resolve()
    allowed_roots = (root, *(path.resolve() for path in trusted_roots))
    within_approved_root = False
    for allowed_root in allowed_roots:
        try:
            resolved.relative_to(allowed_root)
        except ValueError:
            continue
        within_approved_root = True
        break
    if within_approved_root:
        components = [*reversed(lexical.parents), lexical]
        if any(component.is_symlink() for component in components):
            raise ValueError(f"{label} contains a symbolic-link component")
        return lexical
    raise ValueError(f"{label} escapes the approved input roots")


def audit_request_grounding(request: dict[str, Any]) -> dict[str, Any]:
    known_sources = _source_ids(request)
    counts: Counter[str] = Counter()
    missing_source_ids: list[dict[str, Any]] = []
    unclassified: list[dict[str, Any]] = []

    for section in request.get("sections", []):
        for index, block in enumerate(section.get("blocks", []), start=1):
            if block.get("type") in {"page_break", "section_break"}:
                continue
            provenance = block.get("provenance")
            if provenance:
                counts[provenance] += 1
            else:
                unclassified.append({"section": section["heading"], "block": index})
            if provenance == "source":
                block_sources = set(block.get("source_ids", []))
                unknown = sorted(block_sources - known_sources)
                if not block_sources or unknown:
                    missing_source_ids.append(
                        {
                            "section": section["heading"],
                            "block": index,
                            "unknown": unknown,
                            "missing": not block_sources,
                        }
                    )
    return {
        "provenance_counts": dict(sorted(counts.items())),
        "unclassified_blocks": unclassified,
        "source_reference_issues": missing_source_ids,
        "unverified_claims": request.get("unverified_claims", []),
    }


def _display_text(block: dict[str, Any]) -> str:
    text = str(block.get("text", ""))
    if block.get("provenance") == "pending" and not PLACEHOLDER_PATTERNS[
        "pending-confirmation"
    ].search(text):
        return f"[PENDING CONFIRMATION] {text}".strip()
    return text


def _add_text_paragraph(document: DocumentType, block: dict[str, Any]) -> list[Any]:
    block_type = block["type"]
    style = {
        "paragraph": "Normal",
        "quote": "Quote",
        "note": "Note",
        "warning": "Warning",
        "code": "Code",
    }[block_type]
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(_display_text(block))
    return [paragraph._p]


def _add_list(document: DocumentType, block: dict[str, Any]) -> list[Any]:
    style = "List Bullet" if block["type"] == "bullets" else "List Number"
    elements = []
    for item in block["items"]:
        paragraph = document.add_paragraph(str(item), style=style)
        elements.append(paragraph._p)
    return elements


def _add_table(document: DocumentType, block: dict[str, Any]) -> list[Any]:
    headers = [str(value) for value in block["headers"]]
    rows = block["rows"]
    if any(len(row) != len(headers) for row in rows):
        raise ValueError("Every table row must have the same number of cells as headers")
    elements: list[Any] = []
    caption = block.get("caption")
    if caption:
        paragraph = document.add_paragraph(str(caption), style="Caption")
        _set_keep_next(paragraph)
        elements.append(paragraph._p)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_row = table.rows[0]
    _set_repeat_table_header(header_row)
    for index, value in enumerate(headers):
        header_row.cells[index].text = value
        _set_cell_shading(header_row.cells[index], "D9EAF7")
        for run in header_row.cells[index].paragraphs[0].runs:
            run.bold = True
    for values in rows:
        row = table.add_row()
        if not block.get("allow_row_split", False):
            _set_row_cant_split(row)
        for index, value in enumerate(values):
            row.cells[index].text = "" if value is None else str(value)
    elements.append(table._tbl)
    return elements


def _add_figure(
    document: DocumentType,
    block: dict[str, Any],
    *,
    base_dir: Path,
) -> list[Any]:
    image_path = resolve_request_path(
        str(block["path"]), base_dir, label="Figure path"
    )
    if not image_path.is_file():
        raise FileNotFoundError("Figure path does not reference a readable file")

    section = document.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin
    requested_width = Cm(float(block["width_cm"])) if block.get("width_cm") else available_width
    width = min(requested_width, available_width)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=width)
    _set_alt_text(inline_shape, str(block["alt_text"]))
    _set_keep_next(paragraph)
    elements = [paragraph._p]

    if block.get("caption"):
        caption = document.add_paragraph(str(block["caption"]), style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_together = True
        elements.append(caption._p)
    return elements


def add_block(
    document: DocumentType,
    block: dict[str, Any],
    *,
    base_dir: Path,
) -> list[Any]:
    block_type = block["type"]
    if block_type in {"paragraph", "quote", "note", "warning", "code"}:
        return _add_text_paragraph(document, block)
    if block_type in {"bullets", "numbered"}:
        return _add_list(document, block)
    if block_type == "table":
        return _add_table(document, block)
    if block_type == "figure":
        return _add_figure(document, block, base_dir=base_dir)
    if block_type == "page_break":
        paragraph = document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return [paragraph._p]
    if block_type == "section_break":
        section = document.add_section(WD_SECTION.NEW_PAGE)
        if block["orientation"] == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
        return [section._sectPr]
    raise ValueError(f"Unsupported block type: {block_type}")


def add_sections(
    document: DocumentType,
    sections: Iterable[dict[str, Any]],
    *,
    base_dir: Path,
) -> list[Any]:
    elements: list[Any] = []
    for section in sections:
        level = int(section.get("level", 1))
        heading = document.add_heading(str(section["heading"]), level=level)
        _set_keep_next(heading)
        elements.append(heading._p)
        for block in section.get("blocks", []):
            elements.extend(add_block(document, block, base_dir=base_dir))
    return elements


def create_from_request(request: dict[str, Any], request_path: Path) -> DocumentType:
    template_value = request.get("template")
    template = None
    if template_value:
        template = resolve_request_path(
            str(template_value),
            request_path.parent,
            label="Template path",
            trusted_roots=(SKILL_DIR / "assets" / "templates",),
        )
        if not template.is_file():
            raise FileNotFoundError("Template path does not reference a readable file")
    if template:
        with admitted_docx(template) as admitted_template:
            document = Document(str(admitted_template))
    else:
        document = Document()
    if template:
        clear_document_body(document)
    ensure_styles(document, normalize_builtins=not bool(template))

    options = request.get("document", {})
    configure_document(document, options, preserve_template_layout=bool(template))
    properties = {
        "title": request["title"],
        "subject": options.get("subject", request["objective"]),
        "author": options.get("author", ""),
        "keywords": options.get("keywords", ""),
    }
    set_core_properties(document, properties)

    title = document.add_paragraph(style="Title")
    title.add_run(str(request["title"]))
    if request.get("subtitle"):
        subtitle = document.add_paragraph(style="Subtitle")
        subtitle.add_run(str(request["subtitle"]))
    document.add_page_break()

    if options.get("toc"):
        toc_heading = document.add_heading("Contents", level=1)
        _set_keep_next(toc_heading)
        toc = document.add_paragraph()
        _add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Update field to build contents")
        document.add_page_break()

    add_sections(document, request["sections"], base_dir=request_path.parent)
    return document


def paragraph_heading_level(paragraph: Any) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    match = re.fullmatch(r"Heading\s+([1-9])", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _core_properties(document: DocumentType) -> dict[str, Any]:
    properties = document.core_properties
    return {
        "title": properties.title or "",
        "subject": properties.subject or "",
        "author": properties.author or "",
        "last_modified_by": properties.last_modified_by or "",
        "keywords": properties.keywords or "",
        "comments": properties.comments or "",
        "created": properties.created.isoformat() if properties.created else None,
        "modified": properties.modified.isoformat() if properties.modified else None,
    }


def _inspect_admitted_docx(path: Path, *, reported_path: Path) -> dict[str, Any]:
    document = Document(str(path))
    paragraphs = document.paragraphs
    headings = [
        {
            "text": paragraph.text,
            "level": level,
            "style": paragraph.style.name if paragraph.style else "",
        }
        for paragraph in paragraphs
        if (level := paragraph_heading_level(paragraph)) is not None
    ]
    style_usage = Counter(
        paragraph.style.name if paragraph.style else "(none)" for paragraph in paragraphs
    )
    table_header_rows = 0
    table_rows = 0
    for table in document.tables:
        table_rows += len(table.rows)
        if table.rows and table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"):
            table_header_rows += 1

    alt_texts = []
    for shape in document.inline_shapes:
        candidates = shape._inline.xpath(".//wp:docPr")
        alt_texts.append(candidates[0].get("descr", "") if candidates else "")

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        story_part_pattern = re.compile(
            r"word/(?:document|header\d+|footer\d+|footnotes|endnotes|"
            r"comments(?:Extended|Ids|Extensible)?\d*)\.xml"
        )
        story_parts = [
            (name, archive.read(name))
            for name in sorted(names)
            if story_part_pattern.fullmatch(name)
        ]
        story_text: list[str] = []
        comment_personal_metadata: set[str] = set()
        word_elements: Counter[str] = Counter()
        for name, value in story_parts:
            root = ET.fromstring(value)
            for element in root.iter():
                local_name = element.tag.rsplit("}", 1)[-1]
                namespace = (
                    element.tag[1:].split("}", 1)[0]
                    if element.tag.startswith("{")
                    else ""
                )
                if namespace == W_NS:
                    word_elements[local_name] += 1
                if local_name in {"t", "delText", "instrText"} and element.text:
                    story_text.append(element.text)
                for attribute, attribute_value in element.attrib.items():
                    if attribute.rsplit("}", 1)[-1] in {
                        "descr",
                        "name",
                        "title",
                        "tooltip",
                    }:
                        story_text.append(attribute_value)
                if name.startswith("word/comments") and local_name == "comment":
                    for attribute, attribute_value in element.attrib.items():
                        field = attribute.rsplit("}", 1)[-1]
                        if field in {"author", "date", "initials"} and attribute_value:
                            comment_personal_metadata.add(field)
                            if field in {"author", "initials"}:
                                story_text.append(attribute_value)
        all_text = "\n".join(story_text)
        extended_personal_metadata: list[str] = []
        if "docProps/app.xml" in names:
            app_root = ET.fromstring(archive.read("docProps/app.xml"))
            for node in app_root.iter():
                field = node.tag.rsplit("}", 1)[-1]
                if field in {"Company", "HyperlinkBase", "Manager", "Template"}:
                    if (node.text or "").strip():
                        extended_personal_metadata.append(field)
        zip_metadata = {
            "archive_comment": bool(archive.comment),
            "member_comments": sum(bool(info.comment) for info in archive.infolist()),
            "member_extra_fields": sum(bool(info.extra) for info in archive.infolist()),
            "noncanonical_timestamps": sum(
                info.date_time != (1980, 1, 1, 0, 0, 0)
                for info in archive.infolist()
            ),
        }
        relationship_parts = sorted(name for name in names if name.endswith(".rels"))
        package = {
            "parts": len(names),
            "relationship_parts": relationship_parts,
            "comments": word_elements["comment"],
            "insertions": word_elements["ins"],
            "deletions": word_elements["del"],
            "fields": word_elements["instrText"],
            "hyperlinks": word_elements["hyperlink"],
            "bookmarks": word_elements["bookmarkStart"],
            "footnotes_part": "word/footnotes.xml" in names,
            "endnotes_part": "word/endnotes.xml" in names,
            "custom_properties": "docProps/custom.xml" in names,
            "comment_personal_metadata": sorted(comment_personal_metadata),
            "extended_personal_metadata": sorted(extended_personal_metadata),
            "zip_metadata": zip_metadata,
        }

    placeholders = {
        name: len(pattern.findall(all_text))
        for name, pattern in PLACEHOLDER_PATTERNS.items()
        if pattern.search(all_text)
    }
    pii = {
        name: len(pattern.findall(all_text))
        for name, pattern in PII_PATTERNS.items()
        if pattern.search(all_text)
    }

    sections = []
    for section in document.sections:
        sections.append(
            {
                "orientation": (
                    "landscape"
                    if section.orientation == WD_ORIENT.LANDSCAPE
                    else "portrait"
                ),
                "page_width_cm": round(section.page_width.cm, 2),
                "page_height_cm": round(section.page_height.cm, 2),
                "margins_cm": {
                    "top": round(section.top_margin.cm, 2),
                    "right": round(section.right_margin.cm, 2),
                    "bottom": round(section.bottom_margin.cm, 2),
                    "left": round(section.left_margin.cm, 2),
                },
            }
        )

    return {
        "path": str(reported_path),
        "sha256": sha256_file(path),
        "bytes": path.stat().st_size,
        "paragraphs": len(paragraphs),
        "nonempty_paragraphs": sum(bool(paragraph.text.strip()) for paragraph in paragraphs),
        "headings": headings,
        "style_usage": dict(sorted(style_usage.items())),
        "tables": len(document.tables),
        "table_rows": table_rows,
        "tables_with_repeating_header": table_header_rows,
        "inline_figures": len(document.inline_shapes),
        "figures_with_alt_text": sum(bool(value.strip()) for value in alt_texts),
        "sections": sections,
        "headers": [section.header.paragraphs[0].text for section in document.sections],
        "footers": [section.footer.paragraphs[0].text for section in document.sections],
        "core_properties": _core_properties(document),
        "placeholders": placeholders,
        "possible_pii": pii,
        "package": package,
    }


def inspect_docx(path: Path) -> dict[str, Any]:
    """Inspect one stable, admitted snapshot while reporting the caller's path."""
    with admitted_docx(path) as admitted_path:
        return _inspect_admitted_docx(admitted_path, reported_path=path)


def redact_inspection(inspection: dict[str, Any]) -> dict[str, Any]:
    """Return a structural report without document text or metadata values."""
    redacted = dict(inspection)
    redacted["headings"] = [
        {"level": item["level"], "style": item["style"]}
        for item in inspection["headings"]
    ]
    redacted["style_usage"] = {
        "style_count": len(inspection["style_usage"]),
        "paragraphs": sum(inspection["style_usage"].values()),
    }
    redacted["headers"] = {
        "count": len(inspection["headers"]),
        "nonempty": sum(bool(value.strip()) for value in inspection["headers"]),
    }
    redacted["footers"] = {
        "count": len(inspection["footers"]),
        "nonempty": sum(bool(value.strip()) for value in inspection["footers"]),
    }
    redacted["core_properties"] = {
        "populated_fields": sorted(
            key for key, value in inspection["core_properties"].items() if value
        )
    }
    return redacted


def iter_all_paragraphs(document: DocumentType) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
